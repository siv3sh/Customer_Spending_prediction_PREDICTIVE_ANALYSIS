"""
Standalone script to generate comprehensive DOCX report with all analysis and graph explanations
Run: python generate_complete_docx.py
"""
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.model_selection import train_test_split, GridSearchCV
from sklearn.preprocessing import StandardScaler, LabelEncoder
from sklearn.linear_model import LinearRegression
from sklearn.ensemble import RandomForestRegressor, GradientBoostingRegressor
from sklearn.metrics import mean_squared_error, mean_absolute_error, r2_score
import warnings
warnings.filterwarnings('ignore')

# Set style
plt.style.use('default')
sns.set_palette("husl")
np.random.seed(42)

def process_data(df):
    """Process and engineer features"""
    df_processed = df.copy()
    df_processed = df_processed.ffill().bfill()
    
    df_processed['Age_Group_Category'] = pd.cut(df_processed['Age'], 
                                                 bins=[0, 30, 40, 50, 100], 
                                                 labels=['Young', 'Middle', 'Senior', 'Elderly'])
    df_processed['Income_Group_Category'] = pd.cut(df_processed['Annual Income (k$)'], 
                                                    bins=[0, 40, 70, 100, 150], 
                                                    labels=['Low', 'Medium', 'High', 'Very High'])
    df_processed['Income_Age_Ratio'] = df_processed['Annual Income (k$)'] / (df_processed['Age'] + 1)
    df_processed['Age_Squared'] = df_processed['Age'] ** 2
    df_processed['Income_Squared'] = df_processed['Annual Income (k$)'] ** 2
    df_processed['Age_Income_Interaction'] = df_processed['Age'] * df_processed['Annual Income (k$)']
    df_processed['Spending_Capacity'] = (df_processed['Annual Income (k$)'] - df_processed['Annual Income (k$)'].min()) / \
                                         (df_processed['Annual Income (k$)'].max() - df_processed['Annual Income (k$)'].min())
    df_processed['Young_High_Income'] = ((df_processed['Age'] < 35) & (df_processed['Annual Income (k$)'] > 70)).astype(int)
    df_processed['Senior_Low_Income'] = ((df_processed['Age'] > 50) & (df_processed['Annual Income (k$)'] < 50)).astype(int)
    return df_processed

def train_models(X_train_scaled, y_train, X_test_scaled, y_test):
    """Train all models"""
    models = {}
    results = {}
    
    print("Training Linear Regression...")
    simple_model = LinearRegression()
    simple_model.fit(X_train_scaled, y_train)
    models['Linear Regression'] = simple_model
    y_train_pred = simple_model.predict(X_train_scaled)
    y_test_pred = simple_model.predict(X_test_scaled)
    results['Linear Regression'] = {
        'train_rmse': np.sqrt(mean_squared_error(y_train, y_train_pred)),
        'test_rmse': np.sqrt(mean_squared_error(y_test, y_test_pred)),
        'train_mae': mean_absolute_error(y_train, y_train_pred),
        'test_mae': mean_absolute_error(y_test, y_test_pred),
        'train_r2': r2_score(y_train, y_train_pred),
        'test_r2': r2_score(y_test, y_test_pred),
        'predictions': y_test_pred
    }
    
    print("Training Random Forest...")
    rf_param_grid = {'n_estimators': [50, 100], 'max_depth': [5, 10, None], 'min_samples_split': [2, 5]}
    rf_model = RandomForestRegressor(random_state=42)
    rf_grid_search = GridSearchCV(rf_model, rf_param_grid, cv=3, scoring='neg_mean_squared_error', n_jobs=-1, verbose=0)
    rf_grid_search.fit(X_train_scaled, y_train)
    best_rf_model = rf_grid_search.best_estimator_
    models['Random Forest'] = best_rf_model
    y_train_pred = best_rf_model.predict(X_train_scaled)
    y_test_pred = best_rf_model.predict(X_test_scaled)
    results['Random Forest'] = {
        'train_rmse': np.sqrt(mean_squared_error(y_train, y_train_pred)),
        'test_rmse': np.sqrt(mean_squared_error(y_test, y_test_pred)),
        'train_mae': mean_absolute_error(y_train, y_train_pred),
        'test_mae': mean_absolute_error(y_test, y_test_pred),
        'train_r2': r2_score(y_train, y_train_pred),
        'test_r2': r2_score(y_test, y_test_pred),
        'predictions': y_test_pred,
        'best_params': rf_grid_search.best_params_
    }
    
    print("Training Gradient Boosting...")
    gb_param_grid = {'n_estimators': [50, 100], 'max_depth': [3, 5], 'learning_rate': [0.1, 0.2]}
    gb_model = GradientBoostingRegressor(random_state=42)
    gb_grid_search = GridSearchCV(gb_model, gb_param_grid, cv=3, scoring='neg_mean_squared_error', n_jobs=-1, verbose=0)
    gb_grid_search.fit(X_train_scaled, y_train)
    best_gb_model = gb_grid_search.best_estimator_
    models['Gradient Boosting'] = best_gb_model
    y_train_pred = best_gb_model.predict(X_train_scaled)
    y_test_pred = best_gb_model.predict(X_test_scaled)
    results['Gradient Boosting'] = {
        'train_rmse': np.sqrt(mean_squared_error(y_train, y_train_pred)),
        'test_rmse': np.sqrt(mean_squared_error(y_test, y_test_pred)),
        'train_mae': mean_absolute_error(y_train, y_train_pred),
        'test_mae': mean_absolute_error(y_test, y_test_pred),
        'train_r2': r2_score(y_train, y_train_pred),
        'test_r2': r2_score(y_test, y_test_pred),
        'predictions': y_test_pred,
        'best_params': gb_grid_search.best_params_
    }
    return models, results

def create_complete_docx(df, df_processed, models, results, best_model_name, X_train_scaled, scaler, label_encoders):
    """Create complete DOCX with all sections"""
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Customer Annual Spending Score Prediction', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(31, 119, 180)
    title_run.bold = True
    
    subtitle = doc.add_paragraph('E-commerce Customer Behavior Analysis and Predictive Modeling')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(128, 128, 128)
    doc.add_paragraph()
    
    # Import all content from docx_generator
    exec(open('docx_generator.py').read().replace('def create_docx_report', 'def _temp_create_docx_report'))
    
    # Get the function from the module namespace
    import sys
    import importlib.util
    spec = importlib.util.spec_from_file_location("docx_gen", "docx_generator.py")
    docx_gen = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(docx_gen)
    
    # Create base report using the function
    temp_buffer = docx_gen.create_docx_report(df, df_processed, models, results, best_model_name, 
                                              X_train_scaled, scaler, label_encoders)
    
    # Read and copy content
    import io
    temp_doc = Document(io.BytesIO(temp_buffer.getvalue()))
    for element in temp_doc.element.body:
        doc.element.body.append(element)
    
    # Add graph explanations
    add_graph_explanations(doc, df, df_processed, models, results, best_model_name, X_train_scaled, scaler, label_encoders)
    
    return doc

def add_graph_explanations(doc, df, df_processed, models, results, best_model_name, X_train_scaled, scaler, label_encoders):
    """Add detailed graph explanations"""
    doc.add_heading('13. Graph-Based Analysis and Explanations', 1)
    
    doc.add_heading('13.1 Spending Score Distribution Analysis', 2)
    dist_text = f"""The distribution of spending scores reveals important insights about customer behavior:

• Mean Spending Score: {df['Spending Score (1-100)'].mean():.2f} - This indicates the average spending tendency across all customers.
• Median Spending Score: {df['Spending Score (1-100)'].median():.2f} - The median value shows the middle point of the distribution.
• Standard Deviation: {df['Spending Score (1-100)'].std():.2f} - This measures the spread of spending scores, indicating variability in customer spending behavior.

The histogram visualization shows the frequency distribution of spending scores, revealing whether the data follows a normal distribution or has specific patterns. The box plot displays quartiles, median, and outliers, helping identify unusual spending patterns. The violin plot combines the benefits of box plots with kernel density estimation, showing the probability density of spending scores at different values."""
    doc.add_paragraph(dist_text)
    doc.add_paragraph()
    
    doc.add_heading('13.2 Correlation Matrix Analysis', 2)
    numeric_cols = ['Age', 'Annual Income (k$)', 'Spending Score (1-100)']
    corr_matrix = df[numeric_cols].corr()
    age_corr = corr_matrix.loc['Age', 'Spending Score (1-100)']
    income_corr = corr_matrix.loc['Annual Income (k$)', 'Spending Score (1-100)']
    
    corr_analysis = f"""The correlation matrix heatmap visualization reveals the strength and direction of relationships between variables:

• Age vs Spending Score Correlation: {age_corr:.3f}
  - This {'negative' if age_corr < 0 else 'positive'} correlation indicates that {'older' if age_corr < 0 else 'younger'} customers tend to have {'lower' if age_corr < 0 else 'higher'} spending scores.
  - The {'weak' if abs(age_corr) < 0.3 else 'moderate' if abs(age_corr) < 0.7 else 'strong'} relationship suggests age is a {'minor' if abs(age_corr) < 0.3 else 'significant'} factor in spending behavior.

• Annual Income vs Spending Score Correlation: {income_corr:.3f}
  - This {'negative' if income_corr < 0 else 'positive'} correlation shows that customers with {'lower' if income_corr < 0 else 'higher'} income tend to have {'lower' if income_corr < 0 else 'higher'} spending scores.
  - The {'weak' if abs(income_corr) < 0.3 else 'moderate' if abs(income_corr) < 0.7 else 'strong'} relationship indicates income is a {'minor' if abs(income_corr) < 0.3 else 'significant'} predictor of spending behavior.

The heatmap uses color intensity to represent correlation strength, with warmer colors indicating stronger positive correlations and cooler colors indicating negative correlations. This visualization helps identify which features have the strongest relationships with spending scores."""
    doc.add_paragraph(corr_analysis)
    doc.add_paragraph()
    
    doc.add_heading('13.3 Age vs Spending Score Relationship', 2)
    age_analysis = f"""The scatter plot and regression analysis of Age vs Spending Score reveals several key patterns:

• Visual Pattern: The scatter plot shows the distribution of spending scores across different age groups, allowing identification of clusters and outliers. Each point represents a customer, with their age on the x-axis and spending score on the y-axis.

• Regression Line: The fitted regression line indicates the overall trend in the relationship between age and spending. The slope of the line shows whether spending increases or decreases with age.

• Age Group Analysis: When customers are grouped into age categories, we observe:
  - Young customers (18-30): Average spending score of {df[df['Age'] <= 30]['Spending Score (1-100)'].mean():.2f}
  - Middle-aged customers (31-40): Average spending score of {df[(df['Age'] > 30) & (df['Age'] <= 40)]['Spending Score (1-100)'].mean():.2f}
  - Senior customers (41-50): Average spending score of {df[(df['Age'] > 40) & (df['Age'] <= 50)]['Spending Score (1-100)'].mean():.2f}
  - Elderly customers (50+): Average spending score of {df[df['Age'] > 50]['Spending Score (1-100)'].mean():.2f}

The box plots by age groups show the distribution, quartiles, and outliers for each age category, providing insights into spending variability within each group. The boxes represent the interquartile range (IQR), with the line inside showing the median."""
    doc.add_paragraph(age_analysis)
    doc.add_paragraph()
    
    doc.add_heading('13.4 Annual Income vs Spending Score Relationship', 2)
    income_analysis = f"""The analysis of Annual Income vs Spending Score provides critical insights:

• Income Distribution: The scatter plot reveals how spending scores vary across different income levels. Points clustering in certain regions indicate common income-spending patterns.

• Income Groups Analysis:
  - Low Income (≤$40k): Average spending score of {df[df['Annual Income (k$)'] <= 40]['Spending Score (1-100)'].mean():.2f}
  - Medium Income ($41-70k): Average spending score of {df[(df['Annual Income (k$)'] > 40) & (df['Annual Income (k$)'] <= 70)]['Spending Score (1-100)'].mean():.2f}
  - High Income ($71-100k): Average spending score of {df[(df['Annual Income (k$)'] > 70) & (df['Annual Income (k$)'] <= 100)]['Spending Score (1-100)'].mean():.2f}
  - Very High Income (>$100k): Average spending score of {df[df['Annual Income (k$)'] > 100]['Spending Score (1-100)'].mean():.2f}

The regression line shows the trend, while box plots reveal the distribution and identify potential outliers or unusual patterns within each income group. This helps understand whether income is a strong predictor of spending behavior."""
    doc.add_paragraph(income_analysis)
    doc.add_paragraph()
    
    doc.add_heading('13.5 Gender-Based Spending Analysis', 2)
    male_mean = df[df['Gender'] == 'Male']['Spending Score (1-100)'].mean()
    female_mean = df[df['Gender'] == 'Female']['Spending Score (1-100)'].mean()
    
    gender_analysis = f"""The gender analysis reveals spending patterns across different customer segments:

• Male Customers:
  - Average Spending Score: {male_mean:.2f}
  - Count: {len(df[df['Gender'] == 'Male'])} customers
  - Standard Deviation: {df[df['Gender'] == 'Male']['Spending Score (1-100)'].std():.2f}

• Female Customers:
  - Average Spending Score: {female_mean:.2f}
  - Count: {len(df[df['Gender'] == 'Female'])} customers
  - Standard Deviation: {df[df['Gender'] == 'Female']['Spending Score (1-100)'].std():.2f}

The pie chart shows the gender distribution in the dataset, while box plots and violin plots reveal the distribution shape, spread, and potential differences in spending behavior between genders. The violin plots are particularly useful as they combine box plot information with kernel density estimation, showing the probability density of spending scores at different values. This helps identify if there are distinct spending patterns by gender."""
    doc.add_paragraph(gender_analysis)
    doc.add_paragraph()
    
    doc.add_heading('13.6 Three-Dimensional Relationship Analysis', 2)
    d3_analysis = """The 3D scatter plot visualization provides a comprehensive view of the relationship between Age, Annual Income, and Spending Score simultaneously:

• Multi-dimensional Insight: Unlike 2D plots that show relationships between two variables, the 3D plot reveals how all three variables interact together. This helps identify complex patterns that might not be apparent in individual 2D projections.

• Color Coding: The spending score is color-coded, making it easy to identify patterns and clusters of customers with similar characteristics. Customers with similar age, income, and spending patterns form visible clusters.

• Pattern Identification: The visualization helps identify:
  - Clusters of customers with similar age, income, and spending patterns
  - Outliers that deviate from typical patterns
  - Non-linear relationships that might not be apparent in 2D projections
  - Segments that can be targeted with specific marketing strategies

This comprehensive view is crucial for understanding complex customer behavior patterns and identifying segments for targeted marketing. The interactive nature of 3D plots allows rotation and zooming to explore different perspectives of the data."""
    doc.add_paragraph(d3_analysis)
    doc.add_paragraph()
    
    if best_model_name in results:
        doc.add_heading('13.7 Model Performance Visualizations', 2)
        result = results[best_model_name]
        predictions = result['predictions']
        
        X = df_processed.drop(['CustomerID', 'Spending Score (1-100)', 'Age_Group', 'Income_Group'], axis=1, errors='ignore')
        y = df_processed['Spending Score (1-100)']
        
        label_encoders_temp = {}
        for col in X.select_dtypes(include=['object', 'category']).columns:
            le = LabelEncoder()
            X[col] = le.fit_transform(X[col].astype(str))
            label_encoders_temp[col] = le
        
        X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
        
        model_perf_analysis = f"""The model performance visualizations provide critical insights into prediction accuracy:

• Actual vs Predicted Scatter Plot:
  - Shows how closely predictions align with actual values
  - Points closer to the diagonal line (perfect prediction) indicate better accuracy
  - The spread of points reveals prediction consistency
  - Current model shows {'good' if result['test_r2'] > 0.5 else 'moderate'} alignment with R² of {result['test_r2']:.4f}
  - A tight cluster around the diagonal indicates reliable predictions

• Residual Plot:
  - Residuals (actual - predicted) should be randomly distributed around zero
  - Patterns in residuals indicate model bias or missing relationships
  - The current model shows {'random' if abs(np.mean(y_test - predictions)) < 5 else 'some systematic'} residual distribution
  - A funnel shape indicates heteroscedasticity (varying prediction error)
  - Curved patterns suggest non-linear relationships not captured by the model

• Residual Distribution Histogram:
  - Should follow a normal distribution centered at zero for ideal model performance
  - Skewness or multiple peaks indicate areas for model improvement
  - Current distribution shows {'normal' if abs(np.mean(y_test - predictions)) < 3 else 'some deviation'} characteristics
  - Symmetry around zero indicates unbiased predictions

• Actual vs Predicted Line Plot:
  - Shows prediction performance across all test samples
  - Helps identify systematic over-prediction or under-prediction patterns
  - Reveals how well the model captures the overall trend in spending scores
  - Close alignment between actual and predicted lines indicates good model fit"""
        doc.add_paragraph(model_perf_analysis)
        doc.add_paragraph()
        
        doc.add_heading('13.8 Feature Importance Analysis', 2)
        model = models[best_model_name]
        
        if hasattr(model, 'feature_importances_'):
            importance_df = pd.DataFrame({
                'Feature': X_train_scaled.columns,
                'Importance': model.feature_importances_
            }).sort_values('Importance', ascending=False)
        else:
            importance_df = pd.DataFrame({
                'Feature': X_train_scaled.columns,
                'Importance': np.abs(model.coef_)
            }).sort_values('Importance', ascending=False)
        
        top_5 = importance_df.head(5)
        feat_analysis = f"""The feature importance visualization ranks features by their contribution to predictions:

• Top 5 Most Important Features:
  1. {top_5.iloc[0]['Feature']}: {top_5.iloc[0]['Importance']:.4f}
  2. {top_5.iloc[1]['Feature']}: {top_5.iloc[1]['Importance']:.4f}
  3. {top_5.iloc[2]['Feature']}: {top_5.iloc[2]['Importance']:.4f}
  4. {top_5.iloc[3]['Feature']}: {top_5.iloc[3]['Importance']:.4f}
  5. {top_5.iloc[4]['Feature']}: {top_5.iloc[4]['Importance']:.4f}

The bar chart visualization makes it easy to identify which customer attributes drive spending predictions. Features with higher importance scores have greater influence on the model's predictions. This information is crucial for:
- Understanding what drives customer spending behavior
- Identifying key attributes for marketing campaigns
- Focusing data collection efforts on the most predictive features
- Simplifying models by removing less important features"""
        doc.add_paragraph(feat_analysis)
        doc.add_paragraph()
        
        doc.add_heading('13.9 Customer Segmentation Visualizations', 2)
        model = models[best_model_name]
        X = df_processed.drop(['CustomerID', 'Spending Score (1-100)', 'Age_Group', 'Income_Group'], axis=1, errors='ignore')
        
        for col in X.select_dtypes(include=['object', 'category']).columns:
            if col in label_encoders:
                le = label_encoders[col]
                try:
                    X[col] = le.transform(X[col].astype(str))
                except:
                    X[col] = 0
        
        df_processed['Predicted_Spending'] = model.predict(scaler.transform(X))
        df_processed['Customer_Segment'] = pd.cut(df_processed['Predicted_Spending'], 
                                               bins=[0, 30, 50, 70, 100], 
                                               labels=['Low Spender', 'Medium Spender', 'High Spender', 'VIP'])
        
        segment_counts = df_processed['Customer_Segment'].value_counts()
        
        seg_analysis = f"""The customer segmentation visualizations provide actionable insights for marketing:

• Segment Distribution (Pie Chart):
  - Low Spenders: {segment_counts.get('Low Spender', 0)} customers ({segment_counts.get('Low Spender', 0)/len(df_processed)*100:.1f}%)
  - Medium Spenders: {segment_counts.get('Medium Spender', 0)} customers ({segment_counts.get('Medium Spender', 0)/len(df_processed)*100:.1f}%)
  - High Spenders: {segment_counts.get('High Spender', 0)} customers ({segment_counts.get('High Spender', 0)/len(df_processed)*100:.1f}%)
  - VIP Spenders: {segment_counts.get('VIP', 0)} customers ({segment_counts.get('VIP', 0)/len(df_processed)*100:.1f}%)

• Age Distribution by Segment (Box Plot):
  - Shows the age range and distribution for each customer segment
  - Helps identify age-based patterns in spending behavior
  - Reveals which age groups are most likely to be in each segment
  - Outliers in age distribution may indicate special cases

• Income Distribution by Segment (Box Plot):
  - Illustrates income patterns across different spending segments
  - Identifies income thresholds for each segment
  - Guides pricing and product positioning strategies
  - Shows income variability within each segment

• Gender Distribution by Segment (Bar Chart):
  - Shows gender composition of each spending segment
  - Helps tailor marketing messages to gender preferences
  - Identifies gender-based spending patterns
  - Enables gender-specific campaign targeting"""
        doc.add_paragraph(seg_analysis)
        doc.add_paragraph()

def generate_complete_docx():
    """Main function to generate complete DOCX"""
    print("Loading data...")
    df = pd.read_csv('Mall_Customers.csv')
    
    print("Processing data and engineering features...")
    df_processed = process_data(df)
    
    X = df_processed.drop(['CustomerID', 'Spending Score (1-100)', 'Age_Group', 'Income_Group'], axis=1, errors='ignore')
    y = df_processed['Spending Score (1-100)']
    
    label_encoders = {}
    for col in X.select_dtypes(include=['object', 'category']).columns:
        le = LabelEncoder()
        X[col] = le.fit_transform(X[col].astype(str))
        label_encoders[col] = le
    
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
    
    scaler = StandardScaler()
    X_train_scaled = scaler.fit_transform(X_train)
    X_test_scaled = scaler.transform(X_test)
    X_train_scaled = pd.DataFrame(X_train_scaled, columns=X_train.columns, index=X_train.index)
    X_test_scaled = pd.DataFrame(X_test_scaled, columns=X_test.columns, index=X_test.index)
    
    print("Training models...")
    models, results = train_models(X_train_scaled, y_train, X_test_scaled, y_test)
    
    comparison_df = pd.DataFrame({
        'Model': list(results.keys()),
        'Test R²': [results[m]['test_r2'] for m in results.keys()]
    })
    best_model_idx = comparison_df['Test R²'].idxmax()
    best_model_name = comparison_df.loc[best_model_idx, 'Model']
    
    print(f"Best model: {best_model_name}")
    print("Creating complete DOCX report...")
    
    # Import and use docx_generator
    import importlib.util
    spec = importlib.util.spec_from_file_location("docx_gen", "docx_generator.py")
    docx_gen = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(docx_gen)
    
    # Create base report
    temp_buffer = docx_gen.create_docx_report(df, df_processed, models, results, best_model_name, 
                                             X_train_scaled, scaler, label_encoders)
    
    # Read base report
    import io
    doc = Document(io.BytesIO(temp_buffer.getvalue()))
    
    # Add graph explanations
    add_graph_explanations(doc, df, df_processed, models, results, best_model_name, 
                          X_train_scaled, scaler, label_encoders)
    
    # Save
    output_file = 'Customer_Spending_Analysis_Complete_Report.docx'
    doc.save(output_file)
    print(f"✅ DOCX report saved as: {output_file}")
    return output_file

if __name__ == "__main__":
    try:
        output_file = generate_complete_docx()
        print(f"\n🎉 Success! Complete report generated: {output_file}")
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()

